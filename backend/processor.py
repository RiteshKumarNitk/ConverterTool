import os
import json
import logging
from pathlib import Path
import pandas as pd
import pdfplumber
import pytesseract
import cv2
import numpy as np
from docx import Document
from pdf2image import convert_from_path

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ── Tesseract auto-detection ──────────────────────────────────────────────────
_TESSERACT_PATHS = [
    r'C:\Program Files\Tesseract-OCR\tesseract.exe',
    r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
    r'C:\Users\RiteshKumar\AppData\Local\Tesseract-OCR\tesseract.exe',
    r'C:\Users\RITESH\AppData\Local\Tesseract-OCR\tesseract.exe',
]

_tesseract_found = False
for _path in _TESSERACT_PATHS:
    if os.path.exists(_path):
        pytesseract.pytesseract.tesseract_cmd = _path
        logger.info(f"Tesseract found at: {_path}")
        _tesseract_found = True
        break

if not _tesseract_found:
    # Try to find it on PATH (Linux/Mac/Windows with PATH configured)
    import shutil
    _tess_on_path = shutil.which("tesseract")
    if _tess_on_path:
        pytesseract.pytesseract.tesseract_cmd = _tess_on_path
        logger.info(f"Tesseract found on PATH: {_tess_on_path}")
    else:
        logger.warning(
            "Tesseract not found. OCR features will be unavailable. "
            "Install from https://github.com/UB-Mannheim/tesseract/wiki"
        )

# ── Poppler auto-detection ────────────────────────────────────────────────────
_POPPLER_PATHS = [
    r'C:\Program Files\poppler\Library\bin',
    r'C:\Program Files\poppler\bin',
    r'C:\poppler\Library\bin',
    r'C:\poppler\bin',
    r'C:\Program Files\poppler-25.12.0\Library\bin',
    r'C:\Users\RiteshKumar\Downloads\Release-24.02.0-0\poppler-24.02.0\Library\bin',
    r'C:\Users\RITESH\Downloads\Release-24.02.0-0\poppler-24.02.0\Library\bin',
]

_poppler_path_arg = None  # Will be passed to convert_from_path if found
for _p in _POPPLER_PATHS:
    if os.path.exists(_p):
        os.environ["PATH"] += os.pathsep + _p
        _poppler_path_arg = _p
        logger.info(f"Poppler found at: {_p}")
        break

if _poppler_path_arg is None:
    import shutil
    if shutil.which("pdftoppm"):
        logger.info("Poppler found on system PATH.")
    else:
        logger.warning(
            "Poppler not found. PDF→Image OCR fallback will be unavailable. "
            "Install from https://github.com/oschwartz10612/poppler-windows/releases/"
        )

def process_document(file_path: str, output_dir: str, file_id: str):
    """
    Main processing pipeline.
    1. Detect File Type
    2. Extract Text/Tables
    3. Generate Outputs
    """
    file_path_obj = Path(file_path)
    suffix = file_path_obj.suffix.lower()
    
    extracted_data = {
        "text": "",
        "tables": [], # List of list of lists (3D array) or list of dicts
        "metadata": {}
    }

    try:
        if suffix == ".pdf":
            extracted_data = _process_pdf(file_path)
        elif suffix in [".jpg", ".jpeg", ".png", ".bmp", ".tiff"]:
            extracted_data = _process_image(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

        # Save Inputs
        _save_outputs(extracted_data, output_dir, file_id)

        return {
            "status": "completed",
            "file_id": file_id,
            "preview": {
                "text_snippet": extracted_data["text"][:500] if extracted_data["text"] else "",
                "tables_count": len(extracted_data["tables"])
            }
        }

    except Exception as e:
        logger.error(f"Error processing document: {e}")
        raise e

def _process_pdf(file_path):
    text_content = []
    tables_content = []
    
    # Text Extraction
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            # Extract Text
            text = page.extract_text()
            if text:
                text_content.append(text)
            
            # Simple Table Extraction (pdfplumber)
            # For more advanced tables, we'd use Camelot, but it creates dependency hell on Windows often.
            # Using pdfplumber's native table extraction for stability.
            tables = page.extract_tables()
            for table in tables:
                # Clean table: remove Nones
                clean_table = [[cell if cell is not None else "" for cell in row] for row in table]
                tables_content.append(clean_table)

    full_text = "\n".join(text_content)
    
    # If text is empty, it might be a scanned PDF -> OCR
    if not full_text.strip():
        logger.info("No text found in PDF, attempting OCR...")
        full_text, tables_content = _ocr_pdf_images(file_path)

    return {
        "text": full_text,
        "tables": tables_content,
        "type": "pdf"
    }

def _ocr_pdf_images(file_path):
    # Convert PDF to images
    try:
        # Pass poppler_path if we found it, otherwise rely on PATH
        convert_kwargs = {}
        if _poppler_path_arg:
            convert_kwargs["poppler_path"] = _poppler_path_arg

        images = convert_from_path(file_path, **convert_kwargs)
        text_content = []
        
        for img in images:
            # Convert PIL to CV2
            open_cv_image = np.array(img) 
            # OCR
            text = pytesseract.image_to_string(open_cv_image, lang='eng+hin')
            text_content.append(text)
            
        return "\n".join(text_content), []
    except Exception as e:
        logger.warning(f"OCR failed or Poppler not installed: {e}")
        return "[OCR Failed - Please install Poppler and Tesseract]", []

def _process_image(file_path):
    logger.info(f"Processing image: {file_path}")
    
    # Read image
    img = cv2.imread(file_path)
    
    if img is None:
        raise ValueError(f"Failed to load image at {file_path}. The file might be corrupted or the path is invalid.")

    # Preprocessing (Denoise, Grayscale)
    try:
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        
        # Verify Tesseract is ready
        try:
            version = pytesseract.get_tesseract_version()
            logger.info(f"Tesseract Version: {version}")
        except Exception as e:
            raise RuntimeError(
                "Tesseract is not installed or not found. "
                "Please install from https://github.com/UB-Mannheim/tesseract/wiki "
                f"(detail: {e})"
            )

        # OCR
        text = pytesseract.image_to_string(gray, lang='eng+hin')
        logger.info(f"OCR extracted {len(text)} characters.")
        
    except Exception as e:
        logger.error(f"OCR Processing Error: {e}")
        raise e
    
    return {
        "text": text,
        "tables": [],
        "type": "image"
    }

def _save_outputs(data, output_dir, file_id):
    out_path = Path(output_dir)
    
    # 1. JSON
    json_path = out_path / f"{file_id}_result.json"
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
        
    # 2. Excel (Tables)
    if data["tables"]:
        xlsx_path = out_path / f"{file_id}_result.xlsx"
        with pd.ExcelWriter(xlsx_path, engine='openpyxl') as writer:
            for i, table in enumerate(data["tables"]):
                df = pd.DataFrame(table)
                sheet_name = f"Table_{i+1}"
                df.to_excel(writer, sheet_name=sheet_name, index=False, header=False)
    else:
        # Create empty excel if no tables, or just text
        xlsx_path = out_path / f"{file_id}_result.xlsx"
        df = pd.DataFrame({"Content": [data["text"]]})
        df.to_excel(xlsx_path, index=False)

    # 3. Word (Text)
    docx_path = out_path / f"{file_id}_result.docx"
    doc = Document()
    doc.add_heading('Extracted Content', 0)
    doc.add_paragraph(data["text"])
    
    if data["tables"]:
        doc.add_heading('Extracted Tables', level=1)
        for table_data in data["tables"]:
            table = doc.add_table(rows=len(table_data), cols=len(table_data[0]) if table_data else 0)
            table.style = 'Table Grid'
            for i, row in enumerate(table_data):
                row_cells = table.rows[i].cells
                for j, cell_content in enumerate(row):
                    if j < len(row_cells):
                        row_cells[j].text = str(cell_content)
            doc.add_paragraph("\n")
            
    doc.save(docx_path)
